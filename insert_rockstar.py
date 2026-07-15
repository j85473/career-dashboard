import docx
import sys

doc = docx.Document("/Users/JosephLamb/Desktop/Joseph_Lamb_Resume.docx")

# Insert before T-Mobile (which is paragraph 20)
try:
    p_ref = doc.paragraphs[20]
    new_p = p_ref.insert_paragraph_before()
    
    # Copy style from previous bullet (paragraph 19)
    new_p.style = doc.paragraphs[19].style
    
    # Add text
    new_p.add_run("Executed extensive field ride-alongs and co-selling motions with distributor sales reps, actively coaching them on product messaging and competitive positioning to displace rival brands and capture dominant market share.")
    
    doc.save("/Users/JosephLamb/Desktop/Joseph_Lamb_Resume.docx")
    print("Rockstar bullet inserted successfully!")
except Exception as e:
    print(f"Error: {e}")
    sys.exit(1)
