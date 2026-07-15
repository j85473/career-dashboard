import docx

doc = docx.Document("/Users/JosephLamb/Desktop/Joseph_Lamb_Channel_Sales_Resume.docx")

def replace_p(index, new_text):
    p = doc.paragraphs[index]
    if len(p.runs) > 0:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""

replace_p(2, "Channel Sales | Partner Enablement | Ecosystem Growth")
replace_p(3, "Strategic channel sales professional with 6+ years of experience scaling partner ecosystems, driving reseller enablement, and optimizing channel operations across a 4-state territory. Proven track record of managing 7 key strategic partners (spanning retail, B2B hunting, and D2D channels) across 155 locations. Expert at translating go-to-market strategy into field-level execution, engineering automated post-sale workflows, and building profitable relationships with key stakeholders at national accounts.")
replace_p(5, "Channel & Distributor Execution | Partner Enablement | Territory Growth | Go-to-Market (GTM) Strategy | AI-assisted Workflow Development | Post-Sale Support & Retention | Retail Execution | Salesforce, Domo, Zendesk")

replace_p(8, "Field Sales Representative — Channel Sales (AT&T)")
replace_p(9, "Directed diverse Go-To-Market (GTM) channel execution on behalf of AT&T across 7 strategic independent distributors (encompassing Retail, B2B hunting, and Door-to-Door channels), managing a 155-location footprint across a 4-state territory to deliver 15%+ YoY network growth.")
replace_p(10, "Led the regional rollout and field enablement of Sara+ (proprietary order entry and reporting platform); served as the sole implementation resource across the territory, training 7 primary distributor offices and cascading adoption down to the store level.")
replace_p(11, "Architected and deployed a territory-wide distributor retraining program to rebuild core operational execution; standardized the new hire training model, reducing unresolved customer activations from 200+ to under 20 per week within one month.")
replace_p(12, "Engineered an automated post-sale support and retention framework for the region’s largest partner (representing 46% of regional volume), decreasing local account escalations by 82% and driving long-term partner success.")
replace_p(13, "Built direct working relationships with Target and Best Buy Key/National Account Managers, translating real-time store-level issues into field intelligence to accelerate resolution of complex account-level problems.")
replace_p(14, "Managed territory pipeline, distributor performance data, and customer escalation workflows leveraging Salesforce, Domo, and Zendesk.")

replace_p(22, "Managed the company’s largest U.S. territory by coverage area (representing $28M+ in annual volume), embedding directly with distributor reps to deliver a region-leading 10%+ YoY growth across 8 states.")
replace_p(23, "Coordinated distributor product launches, GTM execution, and inventory planning across 3 major markets, supporting 94,000+ annual cases.")
replace_p(24, "Strengthened account-level execution in the Oshkosh market, contributing to 84.56% YoY growth at Woodman’s and 58.30% YoY growth at Kroger/Roundy’s by partnering with distributor reps to improve account coverage and sell-in opportunities (awarded Wisconsin Market of the Year).")

replace_p(27, "Launched the store’s B2B sales program from the ground up through local small business outreach, reaching the top 10% nationally in B2B performance within the first year and generating 33% of the district’s total revenue through business line sales.")
replace_p(28, "Took over a newly opened, unprofitable location in operational disarray; cleared a two-month backlog of unprocessed trade-ins, rebuilt scheduling and coaching structure, stabilized team performance, and turned the store profitable within 90 days.")
replace_p(29, "Created an Excel-based inventory tracking tool connected to handheld scanners, reducing daily phone inventory counts from a lengthy manual process to under three minutes.")

# Delete Barton Associates (indices 15 to 19 inclusive)
for i in range(19, 14, -1):
    p = doc.paragraphs[i]._element
    p.getparent().remove(p)
    p._p = p._element = None

doc.save("/Users/JosephLamb/Desktop/Joseph_Lamb_Channel_Sales_Resume.docx")
print("Resume formatted successfully!")
